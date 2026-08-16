from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "UI_USER_GUIDE.md"
OUTPUT = ROOT / "SRT_Gateway_UI_User_Guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 100, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C8D0DA"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_indent(table, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(indent_dxa))
    indent.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_hyperplain(paragraph, text: str, bold: bool = False) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = bold


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.add_run("SRT Gateway UI User Guide")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Operator manual for creating, monitoring, and troubleshooting routing services")

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    set_table_indent(table)
    set_table_fixed_width(table, [2300, 7060])
    rows = [
        ("UI/API", "http://10.75.51.40:8000"),
        ("Primary video worker", "primary, bound to 10.70.15.3"),
        ("Backup video worker", "backup, bound to 10.71.15.3"),
        ("Monitoring", "Grafana http://10.75.15.3:4000, Prometheus http://10.75.15.3:9090"),
        ("Source", "Generated from UI_USER_GUIDE.md"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        rows.append(cells)
        idx += 1
    if len(rows) >= 2 and all(set(cell.replace(" ", "")) <= {"-", ":"} for cell in rows[1]):
        rows.pop(1)
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    set_table_indent(table)
    widths = [1800] + [int((9360 - 1800) / (cols - 1))] * (cols - 1) if cols > 1 else [9360]
    set_table_fixed_width(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                run = cell.paragraphs[0].add_run(text)
                run.bold = True
            else:
                add_hyperplain(cell.paragraphs[0], text)
    doc.add_paragraph()


def add_bullet_or_number(doc: Document, text: str, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    add_hyperplain(p, text)


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_buffer: list[str] = []
    idx = 0

    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                if code_buffer:
                    add_code_paragraph(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue

        if in_code:
            code_buffer.append(line)
            idx += 1
            continue

        stripped = line.strip()
        if not stripped:
            idx += 1
            continue

        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            if stripped != "# SRT Gateway UI User Guide":
                doc.add_page_break()
                doc.add_heading(stripped[2:], level=1)
            idx += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            idx += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            idx += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
            idx += 1
            continue

        if stripped.startswith("- "):
            add_bullet_or_number(doc, stripped[2:])
            idx += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered_match:
            add_bullet_or_number(doc, numbered_match.group(1), numbered=True)
            idx += 1
            continue

        p = doc.add_paragraph()
        add_hyperplain(p, stripped)
        idx += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("SRT Gateway UI User Guide")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = MUTED

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
